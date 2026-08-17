from __future__ import annotations

import io

from docx import Document
import fitz
from openpyxl import Workbook

from app.parsers import parse_document


def test_markdown_preserves_heading_path_and_table():
    result = parse_document("manual.md", "# VSM-100\n\n## 错误码\n\n| 代码 | 含义 |\n| --- | --- |\n| SYS-NET-042 | 网络 |".encode())
    assert result.status == "ready"
    tables = [block for block in result.blocks if block.block_type == "table"]
    assert len(tables) == 1
    assert "代码 | 含义" in tables[0].text
    assert "SYS-NET-042 | 网络" in tables[0].text
    assert result.blocks[-1].heading_path == ["VSM-100", "错误码"]


def test_docx_parser_preserves_heading_and_table():
    stream = io.BytesIO()
    document = Document()
    document.add_heading("PulseCare", level=1)
    document.add_paragraph("虚构设备资料")
    table = document.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "型号"
    table.rows[0].cells[1].text = "VSM-100"
    document.add_heading("下一节", level=2)
    document.add_paragraph("表格之后的内容")
    document.save(stream)
    result = parse_document("manual.docx", stream.getvalue())
    table_index = next(index for index, block in enumerate(result.blocks) if block.block_type == "table")
    next_heading_index = next(index for index, block in enumerate(result.blocks) if block.text == "下一节")
    assert table_index < next_heading_index
    assert result.blocks[table_index].heading_path == ["PulseCare"]


def test_xlsx_parser_preserves_sheet_and_cell_range():
    stream = io.BytesIO()
    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "兼容矩阵"
    sheet.append(["配件", "型号"])
    sheet.append(["WLM-2", "VSM-100 Pro"])
    workbook.save(stream)
    result = parse_document("matrix.xlsx", stream.getvalue())
    block = result.blocks[0]
    assert block.provenance.sheet == "兼容矩阵"
    assert block.provenance.cell_range == "A1:B1,A2:B2"
    assert block.text == "配件: WLM-2 | 型号: VSM-100 Pro"


def test_xlsx_parser_creates_precise_row_level_blocks_and_splits_tables():
    stream = io.BytesIO()
    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "兼容矩阵"
    sheet.append(["型号", "配件", "最低固件"])
    sheet.append(["VSM-100", "WLM-2", "3.2"])
    sheet.append(["VSM-100 Pro", "WLM-2", "3.4"])
    sheet.append([])
    sheet.append(["错误码", "含义"])
    sheet.append(["SYS-NET-042", "网络未取得地址"])
    workbook.save(stream)

    result = parse_document("matrix.xlsx", stream.getvalue())
    assert len(result.blocks) == 3
    assert result.blocks[1].text == "型号: VSM-100 Pro | 配件: WLM-2 | 最低固件: 3.4"
    assert result.blocks[1].provenance.cell_range == "A1:C1,A3:C3"
    assert result.blocks[2].provenance.cell_range == "A5:B5,A6:B6"


def test_digital_pdf_preserves_page_and_heading():
    document = fitz.open()
    page = document.new_page()
    page.insert_text((72, 72), "VSM-100 Error Codes", fontsize=18)
    page.insert_text((72, 110), "SYS-NET-042 means synthetic network failure.", fontsize=10)
    content = document.tobytes()
    document.close()
    result = parse_document("manual.pdf", content)
    assert result.status == "ready"
    assert all(block.provenance.page == 1 for block in result.blocks)
    assert any(block.block_type == "heading" and "VSM-100" in block.text for block in result.blocks)


def test_pdf_carries_section_heading_to_following_page():
    document = fitz.open()
    first = document.new_page()
    first.insert_text((72, 72), "VSM-100 Network", fontsize=18)
    first.insert_text((72, 110), "First page details.", fontsize=10)
    second = document.new_page()
    second.insert_text((72, 72), "SYS-NET-042 means no valid address.", fontsize=10)
    content = document.tobytes()
    document.close()

    result = parse_document("manual.pdf", content)
    second_page = [block for block in result.blocks if block.provenance.page == 2]
    assert second_page
    assert second_page[0].heading_path == ["VSM-100 Network"]


def test_scanned_pdf_requires_ocr():
    document = fitz.open()
    document.new_page()
    content = document.tobytes()
    document.close()
    result = parse_document("scan.pdf", content)
    assert result.status == "ocr_required"
