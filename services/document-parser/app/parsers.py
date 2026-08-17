from __future__ import annotations

import hashlib
import io
import mimetypes
import re
from pathlib import Path

from bs4 import BeautifulSoup
from docx import Document
from docx.document import Document as _Document
from docx.table import Table
from docx.text.paragraph import Paragraph
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
import fitz
from openpyxl import load_workbook
from openpyxl.utils import get_column_letter

from .models import DocumentBlock, DocumentIR, Provenance


SUPPORTED_SUFFIXES = {".md", ".markdown", ".html", ".htm", ".pdf", ".docx", ".xlsx"}


class UnsupportedDocument(ValueError):
    pass


def parse_document(filename: str, content: bytes, content_type: str = "") -> DocumentIR:
    safe_name = Path(filename or "upload").name
    suffix = Path(safe_name).suffix.lower()
    if suffix not in SUPPORTED_SUFFIXES:
        raise UnsupportedDocument(f"unsupported document type: {suffix or 'unknown'}")
    mime = content_type or mimetypes.guess_type(safe_name)[0] or "application/octet-stream"
    digest = hashlib.sha256(content).hexdigest()
    if suffix in {".md", ".markdown"}:
        blocks = parse_markdown(safe_name, content.decode("utf-8-sig"))
    elif suffix in {".html", ".htm"}:
        blocks = parse_html(safe_name, content.decode("utf-8-sig"))
    elif suffix == ".pdf":
        blocks = parse_pdf(safe_name, content)
    elif suffix == ".docx":
        blocks = parse_docx(safe_name, content)
    else:
        blocks = parse_xlsx(safe_name, content)
    status = "ready"
    warnings: list[str] = []
    if suffix == ".pdf" and not any(block.text.strip() for block in blocks):
        status = "ocr_required"
        warnings.append("PDF contains no extractable text; OCR is required before indexing")
    return DocumentIR(status=status, source_file=safe_name, mime_type=mime, sha256=digest, blocks=blocks, warnings=warnings)


def parse_markdown(filename: str, text: str) -> list[DocumentBlock]:
    blocks: list[DocumentBlock] = []
    headings: list[str] = []
    paragraph: list[str] = []
    table_rows: list[str] = []

    def flush_paragraph() -> None:
        if paragraph:
            value = " ".join(line.strip() for line in paragraph if line.strip())
            if value:
                kind = "list" if all(re.match(r"^\s*(?:[-*+] |\d+[.)] )", line) for line in paragraph) else "paragraph"
                blocks.append(_block(kind, value, filename, headings))
            paragraph.clear()

    def flush_table() -> None:
        if table_rows:
            blocks.append(_block("table", "\n".join(table_rows), filename, headings))
            table_rows.clear()

    for line in text.splitlines():
        heading = re.match(r"^(#{1,6})\s+(.+?)\s*$", line)
        if heading:
            flush_paragraph()
            flush_table()
            level = len(heading.group(1))
            headings[:] = headings[: level - 1]
            headings.append(heading.group(2).strip())
            blocks.append(_block("heading", heading.group(2).strip(), filename, headings))
        elif line.strip().startswith("|") and line.strip().endswith("|"):
            flush_paragraph()
            table_rows.append(line.strip())
        elif not line.strip():
            flush_paragraph()
            flush_table()
        else:
            flush_table()
            paragraph.append(line)
    flush_paragraph()
    flush_table()
    return blocks


def parse_html(filename: str, text: str) -> list[DocumentBlock]:
    soup = BeautifulSoup(text, "html.parser")
    headings: list[str] = []
    blocks: list[DocumentBlock] = []
    for element in soup.find_all(["h1", "h2", "h3", "h4", "h5", "h6", "p", "li", "pre", "table"]):
        value = element.get_text(" | " if element.name == "table" else " ", strip=True)
        if not value:
            continue
        if element.name.startswith("h"):
            level = int(element.name[1])
            headings[:] = headings[: level - 1]
            headings.append(value)
            kind = "heading"
        elif element.name == "table":
            kind = "table"
        elif element.name == "li":
            kind = "list"
        elif element.name == "pre":
            kind = "code"
        else:
            kind = "paragraph"
        blocks.append(_block(kind, value, filename, headings))
    return blocks


def parse_pdf(filename: str, content: bytes) -> list[DocumentBlock]:
    blocks: list[DocumentBlock] = []
    document = fitz.open(stream=content, filetype="pdf")
    headings: list[str] = []
    try:
        for page_number, page in enumerate(document, start=1):
            table_boxes: list[fitz.Rect] = []
            candidates: list[tuple[float, float, str, object]] = []
            try:
                finder = page.find_tables()
                for table in finder.tables:
                    rows = table.extract()
                    value = "\n".join(" | ".join("" if cell is None else str(cell).strip() for cell in row) for row in rows)
                    if value.strip(" |\n"):
                        bbox = fitz.Rect(table.bbox)
                        table_boxes.append(bbox)
                        candidates.append((bbox.y0, bbox.x0, "table", value))
            except Exception:
                # Some PDFs do not expose table geometry; text extraction still
                # provides a safe, page-addressable fallback.
                table_boxes = []
            page_dict = page.get_text("dict", sort=True)
            for raw in page_dict.get("blocks", []):
                if raw.get("type") != 0:
                    continue
                bbox = fitz.Rect(raw.get("bbox", (0, 0, 0, 0)))
                if any(bbox.intersects(table_box) and bbox.get_area() > 0 for table_box in table_boxes):
                    continue
                spans = [span for line in raw.get("lines", []) for span in line.get("spans", [])]
                value = "\n".join("".join(str(span.get("text", "")) for span in line.get("spans", [])).strip() for line in raw.get("lines", [])).strip()
                if not value:
                    continue
                max_size = max((float(span.get("size", 0)) for span in spans), default=0)
                bold = any(int(span.get("flags", 0)) & 16 for span in spans)
                candidates.append((bbox.y0, bbox.x0, "text", (value, max_size, bold)))

            # PyMuPDF exposes text and tables through separate APIs. Sorting
            # them back into visual order is essential: otherwise every table
            # loses the section heading that gives its columns meaning.
            for _, _, kind, payload in sorted(candidates, key=lambda item: (item[0], item[1])):
                if kind == "table":
                    blocks.append(DocumentBlock(
                        block_type="table", text=str(payload), heading_path=list(headings),
                        provenance=Provenance(source_file=filename, page=page_number),
                    ))
                    continue
                value, max_size, bold = payload  # type: ignore[misc]
                is_heading = len(value) <= 120 and (max_size >= 14 or bold)
                if is_heading:
                    headings = [value]
                blocks.append(DocumentBlock(
                    block_type="heading" if is_heading else "paragraph",
                    text=value, heading_path=list(headings),
                    provenance=Provenance(source_file=filename, page=page_number),
                ))
    finally:
        document.close()
    return blocks


def parse_docx(filename: str, content: bytes) -> list[DocumentBlock]:
    document = Document(io.BytesIO(content))
    blocks: list[DocumentBlock] = []
    headings: list[str] = []
    for item in _iter_docx_blocks(document):
        if isinstance(item, Paragraph):
            value = item.text.strip()
            if not value:
                continue
            style = item.style.name if item.style else ""
            match = re.match(r"Heading\s+(\d+)", style, re.IGNORECASE)
            if match:
                level = int(match.group(1))
                headings[:] = headings[: level - 1]
                headings.append(value)
                kind = "heading"
            else:
                kind = "list" if "List" in style else "paragraph"
            blocks.append(_block(kind, value, filename, headings))
            continue
        rows = [" | ".join(cell.text.strip() for cell in row.cells) for row in item.rows]
        value = "\n".join(row for row in rows if row.strip(" |"))
        if value:
            blocks.append(_block("table", value, filename, headings))
    return blocks


def parse_xlsx(filename: str, content: bytes) -> list[DocumentBlock]:
    workbook = load_workbook(io.BytesIO(content), read_only=True, data_only=True)
    blocks: list[DocumentBlock] = []
    try:
        for sheet in workbook.worksheets:
            segments: list[list[tuple[int, list[tuple[int, str]]]]] = []
            current: list[tuple[int, list[tuple[int, str]]]] = []
            previous_row = 0
            for row_number, row in enumerate(sheet.iter_rows(), start=1):
                values = [(column, "" if cell.value is None else str(cell.value).strip()) for column, cell in enumerate(row, start=1)]
                nonempty = [(column, value) for column, value in values if value]
                if not nonempty:
                    if current:
                        segments.append(current)
                        current = []
                    continue
                if current and row_number > previous_row + 1:
                    segments.append(current)
                    current = []
                current.append((row_number, nonempty))
                previous_row = row_number
            if current:
                segments.append(current)

            for segment in segments:
                _append_xlsx_segment(blocks, filename, sheet.title, segment)
    finally:
        workbook.close()
    return blocks


def _iter_docx_blocks(document: _Document):
    """Yield paragraphs and tables in their actual body order."""
    for child in document.element.body.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, document)
        elif isinstance(child, CT_Tbl):
            yield Table(child, document)


def _append_xlsx_segment(
    blocks: list[DocumentBlock],
    filename: str,
    sheet_name: str,
    rows: list[tuple[int, list[tuple[int, str]]]],
) -> None:
    """Create row-addressable table blocks while retaining header semantics."""
    header_row, header_cells = rows[0]
    headers = {column: value for column, value in header_cells}
    header_range = _xlsx_row_range(header_row, header_cells)
    if len(rows) == 1:
        blocks.append(DocumentBlock(
            block_type="table",
            text=" | ".join(value for _, value in header_cells),
            heading_path=[sheet_name],
            provenance=Provenance(source_file=filename, sheet=sheet_name, cell_range=header_range),
        ))
        return

    for row_number, cells in rows[1:]:
        values = []
        for column, value in cells:
            label = headers.get(column) or get_column_letter(column)
            values.append(f"{label}: {value}")
        row_range = _xlsx_row_range(row_number, cells)
        blocks.append(DocumentBlock(
            block_type="table",
            text=" | ".join(values),
            heading_path=[sheet_name],
            provenance=Provenance(
                source_file=filename,
                sheet=sheet_name,
                cell_range=f"{header_range},{row_range}",
            ),
        ))


def _xlsx_row_range(row_number: int, cells: list[tuple[int, str]]) -> str:
    start = get_column_letter(min(column for column, _ in cells))
    end = get_column_letter(max(column for column, _ in cells))
    return f"{start}{row_number}:{end}{row_number}"


def _block(kind: str, text: str, filename: str, headings: list[str]) -> DocumentBlock:
    return DocumentBlock(block_type=kind, text=text, heading_path=list(headings), provenance=Provenance(source_file=filename))
