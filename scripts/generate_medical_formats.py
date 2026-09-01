#!/usr/bin/env python3
"""Generate fictional multi-format fixtures from the medical-device corpus."""

from __future__ import annotations

from datetime import datetime
from pathlib import Path
import re
import zipfile

import fitz
from docx import Document
from openpyxl import Workbook


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "datasets/domains/medical-device/corpus/documents"
OUTPUT = ROOT / "data/medical/generated"
FIXED_TIME = datetime(2026, 1, 1, 0, 0, 0)


def normalize_zip(path: Path) -> None:
    """Remove wall-clock ZIP metadata so generated Office files are stable."""
    temporary = path.with_suffix(path.suffix + ".tmp")
    with zipfile.ZipFile(path, "r") as source, zipfile.ZipFile(
        temporary, "w", compression=zipfile.ZIP_DEFLATED, compresslevel=9
    ) as target:
        for original in source.infolist():
            info = zipfile.ZipInfo(original.filename, (2026, 1, 1, 0, 0, 0))
            info.compress_type = original.compress_type
            info.create_system = original.create_system
            info.external_attr = original.external_attr
            info.flag_bits = original.flag_bits
            content = source.read(original.filename)
            # openpyxl overwrites the workbook's modified property with the
            # wall clock during save(), even when callers set it explicitly.
            # That made every bootstrap look like a new document revision and
            # triggered needless parsing, embedding calls and Milvus writes.
            if original.filename == "docProps/core.xml":
                content = re.sub(
                    rb"(<dcterms:modified\b[^>]*>).*?(</dcterms:modified>)",
                    rb"\g<1>2026-01-01T00:00:00Z\g<2>",
                    content,
                    flags=re.DOTALL,
                )
            target.writestr(info, content)
    temporary.replace(path)


def docx_fixture() -> None:
    document = Document()
    document.core_properties.created = FIXED_TIME
    document.core_properties.modified = FIXED_TIME
    document.add_heading("VSM-100 软件 2.6 系统错误码（DOCX 测试副本）", level=1)
    document.add_paragraph("本文件由项目脚本从虚构测试资料派生，仅用于解析与跨格式检索验证。")
    document.add_heading("SYS-NET-042", level=2)
    document.add_paragraph("VSM-100 软件 2.6 中表示设备网络接口无法取得有效地址。")
    table = document.add_table(rows=1, cols=3)
    for cell, value in zip(table.rows[0].cells, ("错误码", "适用型号", "排查入口"), strict=True):
        cell.text = value
    row = table.add_row().cells
    for cell, value in zip(row, ("SYS-NET-042", "VSM-100", "设置 > 网络 > 接口状态"), strict=True):
        cell.text = value
    document.save(OUTPUT / "vsm100-error-codes-fw2.6.docx")
    normalize_zip(OUTPUT / "vsm100-error-codes-fw2.6.docx")

    revision = Document()
    revision.core_properties.created = FIXED_TIME
    revision.core_properties.modified = FIXED_TIME
    revision.add_heading("VSM-100 软件 2.6 系统错误码（修订演示 R2）", level=1)
    revision.add_paragraph(
        "本文件是同一虚构文档的第二修订，仅用于验证版本差异、增量索引和检索命中。"
    )
    revision.add_heading("SYS-NET-042", level=2)
    revision.add_paragraph(
        "VSM-100 软件 2.6 中表示设备网络接口无法取得有效 IPv4 地址；R2 补充要求先确认物理链路状态。"
    )
    revision.add_heading("排查前置条件", level=2)
    revision.add_paragraph("确认网线已连接，接口指示灯常亮或闪烁，并记录当前 DHCP 状态。")
    revision.add_paragraph("若使用静态地址，核对地址、子网掩码与网关位于同一规划网段。")
    table = revision.add_table(rows=1, cols=4)
    for cell, value in zip(
        table.rows[0].cells,
        ("错误码", "适用型号", "排查入口", "升级条件"),
        strict=True,
    ):
        cell.text = value
    row = table.add_row().cells
    for cell, value in zip(
        row,
        (
            "SYS-NET-042",
            "VSM-100",
            "设置 > 网络 > 接口状态",
            "链路正常且重新获取地址失败 3 次后联系服务人员",
        ),
        strict=True,
    ):
        cell.text = value
    revision.save(OUTPUT / "vsm100-error-codes-fw2.6-r2.docx")
    normalize_zip(OUTPUT / "vsm100-error-codes-fw2.6-r2.docx")


def xlsx_fixture() -> None:
    workbook = Workbook()
    workbook.properties.created = FIXED_TIME
    workbook.properties.modified = FIXED_TIME
    sheet = workbook.active
    sheet.title = "兼容矩阵"
    sheet.append(["型号", "软件版本", "配件", "最低固件", "结论"])
    sheet.append(["VSM-100", "2.6", "HUB-4", "1.8", "兼容"])
    sheet.append(["VSM-100", "2.6", "WLM-2", "3.2", "兼容"])
    sheet.append(["VSM-100 Pro", "2.6", "WLM-2", "3.4", "兼容；不得沿用 VSM-100 的 3.2 结论"])
    note = workbook.create_sheet("说明")
    note.append(["范围", "完全虚构的跨格式检索测试数据"])
    workbook.save(OUTPUT / "vsm100-compatibility-fw2.6.xlsx")
    normalize_zip(OUTPUT / "vsm100-compatibility-fw2.6.xlsx")


def pdf_fixture() -> None:
    document = fitz.open()
    page = document.new_page()
    page.insert_text((60, 72), "PulseCare Field Correction Test Notice FC-2026-04", fontsize=15)
    lines = [
        "Synthetic data only. Not for real device operation.",
        "Affected model: VSM-100 Pro",
        "Affected software: 2.5.0 through 2.5.3",
        "Affected lots: L26A01 through L26A07",
        "Target software after correction: 2.5.4",
        "The applicability result must be computed by the deterministic policy tool.",
    ]
    for index, line in enumerate(lines):
        page.insert_text((60, 112 + index * 24), line, fontsize=11)
    document.set_metadata({
        "title": "FC-2026-04 synthetic PDF fixture",
        "creationDate": "D:20260101000000+08'00'",
        "modDate": "D:20260101000000+08'00'",
    })
    document.save(
        OUTPUT / "field-correction-fc-2026-04.pdf",
        garbage=4,
        deflate=True,
        no_new_id=True,
    )


def html_fixture() -> None:
    body = (SOURCE / "vsm100-release-notes-fw2.6.md").read_text(encoding="utf-8")
    escaped = body.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
    html = f"""<!doctype html><html lang=\"zh-CN\"><head><meta charset=\"utf-8\"><title>VSM-100 软件 2.6 发布说明</title></head><body><h1>VSM-100 软件 2.6 发布说明（HTML 测试副本）</h1><p>完全虚构，仅用于解析验证。</p><pre>{escaped}</pre></body></html>"""
    (OUTPUT / "vsm100-release-notes-fw2.6.html").write_text(html, encoding="utf-8")


def main() -> None:
    OUTPUT.mkdir(parents=True, exist_ok=True)
    docx_fixture()
    xlsx_fixture()
    pdf_fixture()
    html_fixture()
    for path in sorted(OUTPUT.iterdir()):
        print(path.relative_to(ROOT))


if __name__ == "__main__":
    main()
