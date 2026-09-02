#!/usr/bin/env python3
"""Export real parser/cleaner/chunker artifacts through the authenticated API.

The fixtures are synthetic and are sent to the no-index evaluation endpoint.
No source file, access token, Document IR or chunk is persisted by RagLab.
"""

from __future__ import annotations

import argparse
import io
import json
import os
import tempfile
from pathlib import Path

import httpx
import pymupdf
from openpyxl import Workbook
from docx import Document


DEVELOPMENT_CASES = (
    ("dev-ocr-aed-critical-001", "synthetic-aed-troubleshooting-r1", "aed-scan.pdf", "application/pdf"),
    ("dev-cleaning-margin-002", "synthetic-vsm100-operator-r1", "vsm100-margins.pdf", "application/pdf"),
    ("dev-table-version-003", "synthetic-vsm-compatibility-r1", "vsm-compatibility.xlsx", "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"),
    ("dev-chunk-long-procedure-004", "synthetic-long-service-procedure-r1", "long-procedure.md", "text/markdown"),
    ("dev-ocr-degraded-fallback-005", "synthetic-monitor-degraded-ocr-r1", "monitor-degraded-development.pdf", "application/pdf"),
    ("dev-version-scope-filter-006", "synthetic-vsm460-network-r4", "vsm460-network-current.md", "text/markdown"),
    ("dev-version-scope-filter-006", "synthetic-vsm460-network-r3", "vsm460-network-history.md", "text/markdown"),
    ("dev-model-suffix-filter-007", "synthetic-vsm420-pro-power-r1", "vsm420-pro-power.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document"),
    ("dev-model-suffix-filter-007", "synthetic-vsm420-power-r1", "vsm420-power.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document"),
    ("dev-lot-scope-filter-008", "synthetic-c5-field-notice-r2", "c5-field-notice-current.html", "text/html"),
    ("dev-lot-scope-filter-008", "synthetic-c5-field-notice-r1", "c5-field-notice-history.html", "text/html"),
)

HOLDOUT_CASES = (
    ("holdout-version-conflict-004", "synthetic-vsm450-network-r3", "vsm450-network-current.md", "text/markdown"),
    ("holdout-version-conflict-004", "synthetic-vsm450-network-r2", "vsm450-network-history.md", "text/markdown"),
    ("holdout-similar-model-005", "synthetic-vsm410-pro-power-r1", "vsm410-pro-power.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document"),
    ("holdout-similar-model-005", "synthetic-vsm410-power-r1", "vsm410-power.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document"),
    ("holdout-superseded-notice-006", "synthetic-c3-field-notice-r3", "c3-field-notice-current.html", "text/html"),
    ("holdout-superseded-notice-006", "synthetic-c3-field-notice-r1", "c3-field-notice-obsolete.html", "text/html"),
)

DOCUMENT_METADATA = {
    "synthetic-aed-troubleshooting-r1": {"model_codes": ["BeneHeart C2"], "document_revision": "R1", "authority_level": "service_manual"},
    "synthetic-vsm100-operator-r1": {"model_codes": ["VSM-100"], "document_revision": "R1", "authority_level": "operator_manual"},
    "synthetic-vsm-compatibility-r1": {"model_codes": ["VSM-100", "VSM-100 Pro"], "software_version_from": "2.6", "software_version_to": "2.6", "document_revision": "R1"},
    "synthetic-long-service-procedure-r1": {"model_codes": ["VSM-300"], "document_revision": "R1", "authority_level": "service_manual"},
    "synthetic-monitor-degraded-ocr-r1": {"model_codes": ["BeneVision N12"], "document_revision": "R1"},
    "synthetic-vsm450-network-r3": {"model_codes": ["VSM-450"], "software_version_from": "4.2", "software_version_to": "4.2", "document_revision": "R3", "supersedes": ["synthetic-vsm450-network-r2"], "authority_level": "current_service_manual"},
    "synthetic-vsm450-network-r2": {"model_codes": ["VSM-450"], "software_version_from": "3.8", "software_version_to": "3.8", "document_revision": "R2", "authority_level": "superseded"},
    "synthetic-vsm410-pro-power-r1": {"model_codes": ["VSM-410 Pro"], "document_revision": "R1", "authority_level": "service_manual"},
    "synthetic-vsm410-power-r1": {"model_codes": ["VSM-410"], "document_revision": "R1", "authority_level": "service_manual"},
    "synthetic-c3-field-notice-r3": {"model_codes": ["BeneHeart C3"], "document_revision": "R3", "supersedes": ["synthetic-c3-field-notice-r1"], "affected_lots": ["LOT-K2608"], "authority_level": "current_field_notice"},
    "synthetic-c3-field-notice-r1": {"model_codes": ["BeneHeart C3"], "document_revision": "R1", "affected_lots": ["LOT-K2501"], "authority_level": "superseded"},
    "synthetic-vsm460-network-r4": {"model_codes": ["VSM-460"], "software_version_from": "5.1", "software_version_to": "5.1", "document_revision": "R4", "supersedes": ["synthetic-vsm460-network-r3"], "authority_level": "current_service_manual"},
    "synthetic-vsm460-network-r3": {"model_codes": ["VSM-460"], "software_version_from": "4.7", "software_version_to": "4.7", "document_revision": "R3", "authority_level": "superseded"},
    "synthetic-vsm420-pro-power-r1": {"model_codes": ["VSM-420 Pro"], "document_revision": "R1", "authority_level": "service_manual"},
    "synthetic-vsm420-power-r1": {"model_codes": ["VSM-420"], "document_revision": "R1", "authority_level": "service_manual"},
    "synthetic-c5-field-notice-r2": {"model_codes": ["BeneHeart C5"], "document_revision": "R2", "supersedes": ["synthetic-c5-field-notice-r1"], "affected_lots": ["LOT-M2701"], "authority_level": "current_field_notice"},
    "synthetic-c5-field-notice-r1": {"model_codes": ["BeneHeart C5"], "document_revision": "R1", "affected_lots": ["LOT-M2602"], "authority_level": "superseded"},
}

LONG_PROCEDURE = (
    "处理步骤：先断开 VSM-300 外部电源并记录 BAT-SVC-088，再确认备用电池已隔离；"
    "随后依次检查电池连接器、保险丝状态与电源管理板指示灯，禁止跳过绝缘检查；"
    "完成后恢复供电并观察十五分钟，确认错误码不再出现，最后把测量值、设备序列号、软件版本和复核人员写入服务记录。"
)

VSM450_CURRENT_PROCEDURE = (
    "NET-LINK-204 当前处理流程：仅适用于 VSM-450 软件 4.2。先在维护界面导出网络诊断包并记录设备序列号、"
    "网卡地址和告警时间，不得直接恢复出厂设置；随后核对有线接口指示灯、交换机端口 VLAN 与网关地址，"
    "确认现场没有把 VSM-450 历史版本 3.8 的固定地址模板复制到 4.2。保持患者监护功能由备用设备接管后，"
    "依次断开数据上传服务、刷新网络证书、重新加载 4.2 配置，再执行网关连通性和时间同步测试。若证书链"
    "校验失败，只允许安装工单中批准的当前证书，不得回退到旧版共享证书。恢复连接后连续观察十五分钟，"
    "确认 NET-LINK-204 不再出现，并分别验证中央站、归档服务器和远程维护通道。最后把诊断包哈希、变更前后"
    "地址、软件版本 4.2、测试结果和复核人员写入服务记录；任一验证失败都必须恢复变更前配置并升级厂商支持。"
)


def cjk_font() -> str:
    configured = os.getenv("RAGLAB_CJK_FONT", "").strip()
    candidates = [
        configured,
        "/System/Library/Fonts/PingFang.ttc",
        "/System/Library/Fonts/STHeiti Light.ttc",
        "/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc",
        "/usr/share/fonts/truetype/wqy/wqy-zenhei.ttc",
    ]
    for candidate in candidates:
        if candidate and Path(candidate).is_file():
            return candidate
    raise RuntimeError("a CJK font is required; set RAGLAB_CJK_FONT to a local .ttf/.ttc file")


def add_text(
    page: pymupdf.Page,
    position: tuple[float, float],
    value: str,
    size: float,
    font_file: str,
    *,
    color: tuple[float, float, float] = (0, 0, 0),
) -> None:
    page.insert_font(fontname="raglab-cjk", fontfile=font_file)
    page.insert_text(position, value, fontsize=size, fontname="raglab-cjk", color=color)


def scanned_aed(font_file: str) -> bytes:
    source = pymupdf.open()
    page = source.new_page(width=800, height=1100)
    add_text(page, (70, 100), "AED 设备故障排查", 28, font_file)
    # Keep body typography consistent. A mixed 20/18pt layout repeatedly
    # caused PP-DocLayout-S to omit the fourth line even though OCR itself
    # recognized the identifiers; the controlled probe is retained as a real
    # layout-model Bad Case in the evaluation report.
    add_text(page, (70, 160), "型号：BeneHeart C2", 18, font_file)
    add_text(page, (70, 220), "错误码：BAT-LOW-021", 18, font_file)
    add_text(page, (70, 280), "BAT-LOW-021 处理：连接交流电源并检查电池状态", 18, font_file)
    add_text(page, (70, 340), "仅授权服务人员执行", 18, font_file)
    pixmap = page.get_pixmap(matrix=pymupdf.Matrix(2, 2), alpha=False)
    source.close()
    scanned = pymupdf.open()
    image_page = scanned.new_page(width=pixmap.width, height=pixmap.height)
    image_page.insert_image(image_page.rect, pixmap=pixmap)
    content = scanned.tobytes(deflate=True)
    scanned.close()
    return content


def repeated_margin_pdf(font_file: str) -> bytes:
    document = pymupdf.open()
    bodies = (
        "VSM-100 网络配置说明",
        "SYS-NET-042 请检查网络接口和网关配置",
        "执行变更前记录当前软件版本，禁止删除相似型号 VSM-100 Pro 的独立说明。",
    )
    for index, body in enumerate(bodies, start=1):
        page = document.new_page(width=595, height=842)
        add_text(page, (60, 45), "PulseCare Medical Devices", 10, font_file)
        add_text(page, (60, 160), body, 15, font_file)
        add_text(page, (270, 815), f"第 {index} 页", 10, font_file)
    content = document.tobytes(deflate=True)
    document.close()
    return content


def compatibility_xlsx(target: Path) -> bytes:
    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "兼容矩阵"
    sheet.append(["型号", "软件版本", "硬件修订"])
    sheet.append(["VSM-100", "2.6", "HR-1"])
    sheet.append(["VSM-100 Pro", "2.6", "HR-2"])
    # Enough nearby rows to make embedding amplification visible when overlap
    # changes, while retaining exact model/version strings.
    for index in range(1, 26):
        sheet.append([f"VSM-200-{index:02d}", "3.1", f"HX-{index:02d}"])
    workbook.save(target)
    return target.read_bytes()


def long_procedure_markdown() -> bytes:
    filler_prefix = ("背景说明用于模拟真实服务手册中的连续长段落，当前句不包含操作结论。" * 10)[:270]
    filler_suffix = ("补充说明用于模拟审计、备件和交接要求，不能替代正式服务流程。" * 12)[:320]
    return f"# 长段落切分验证\n\n{filler_prefix}{LONG_PROCEDURE}{filler_suffix}\n".encode()


def rotated_pump_scan(font_file: str) -> bytes:
    source = pymupdf.open()
    page = source.new_page(width=800, height=1100)
    add_text(page, (70, 110), "BeneFusion uVP 输注泵服务手册", 25, font_file)
    add_text(page, (70, 190), "错误码 INF-OCC-017", 19, font_file)
    add_text(page, (70, 260), "INF-OCC-017 停止输注并检查管路", 19, font_file)
    add_text(page, (70, 330), "确认夹闭、折弯或堵塞已解除后再按规程复核。", 17, font_file)
    pixmap = page.get_pixmap(matrix=pymupdf.Matrix(2, 2), alpha=False)
    source.close()
    scanned = pymupdf.open()
    image_page = scanned.new_page(width=pixmap.width, height=pixmap.height)
    image_page.insert_image(image_page.rect, pixmap=pixmap)
    # Preserve a real 90-degree page-orientation probe. The OCR worker, rather
    # than the fixture generator, must prove that it can normalize the page.
    image_page.set_rotation(90)
    content = scanned.tobytes(deflate=True)
    scanned.close()
    return content


def low_contrast_monitor_scan(font_file: str) -> bytes:
    source = pymupdf.open()
    page = source.new_page(width=595, height=842)
    ink = (0.72, 0.72, 0.72)
    add_text(page, (55, 105), "BeneVision N1 现场服务记录", 20, font_file, color=ink)
    add_text(page, (55, 175), "设备序列号：N1-SYNTHETIC-204", 13, font_file, color=ink)
    add_text(page, (55, 230), "低对比度扫描件必须由人工确认后发布。", 13, font_file, color=ink)
    # 150 DPI is intentionally close to a common archive-scan boundary. No
    # expected status is hard-coded here; the parser confidence decides it.
    pixmap = page.get_pixmap(dpi=150, alpha=False)
    source.close()
    scanned = pymupdf.open()
    image_page = scanned.new_page(width=pixmap.width, height=pixmap.height)
    image_page.insert_image(image_page.rect, pixmap=pixmap)
    content = scanned.tobytes(deflate=True)
    scanned.close()
    return content


def degraded_monitor_development_scan(font_file: str) -> bytes:
    source = pymupdf.open()
    page = source.new_page(width=595, height=842)
    ink = (0.70, 0.70, 0.70)
    add_text(page, (55, 105), "BeneVision N12 维修接收单", 20, font_file, color=ink)
    add_text(page, (55, 175), "设备编号：N12-DEVELOPMENT-031", 13, font_file, color=ink)
    add_text(page, (55, 230), "结构识别降级结果必须人工复核。", 13, font_file, color=ink)
    pixmap = page.get_pixmap(dpi=180, alpha=False)
    source.close()
    scanned = pymupdf.open()
    image_page = scanned.new_page(width=pixmap.width, height=pixmap.height)
    image_page.insert_image(image_page.rect, pixmap=pixmap)
    content = scanned.tobytes(deflate=True)
    scanned.close()
    return content


def aed_maintenance_docx() -> bytes:
    document = Document()
    document.add_heading("维护前检查", level=1)
    document.add_paragraph("确认设备已退出临床使用", style="List Bullet")
    document.add_paragraph("记录设备型号、软件版本和维护工单号", style="List Bullet")
    table = document.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "型号"
    table.rows[0].cells[1].text = "软件版本"
    row = table.add_row().cells
    row[0].text = "C2"
    row[1].text = "2.6"
    document.add_paragraph("MAINT-003 仅授权服务人员执行")
    output = io.BytesIO()
    document.save(output)
    return output.getvalue()


def vsm450_network_current_markdown() -> bytes:
    return f"# VSM-450 4.2 网络恢复流程\n\n{VSM450_CURRENT_PROCEDURE}\n".encode()


def vsm450_network_history_markdown() -> bytes:
    return (
        "# VSM-450 3.8 历史网络说明\n\n"
        "NET-LINK-204 历史处理流程：仅适用于 VSM-450 软件 3.8。该版本使用固定地址模板，"
        "完成接口检查后重新导入 3.8 备份。此修订已被 4.2 网络恢复流程取代，不得用于软件 4.2。\n"
    ).encode()


def power_docx(model: str, document_revision: str, action: str) -> bytes:
    document = Document()
    document.add_heading(f"{model} 电源告警说明", level=1)
    document.add_paragraph(f"文档修订：{document_revision}")
    document.add_paragraph(f"型号：{model}")
    document.add_paragraph(f"PWR-017 处理：{action}")
    output = io.BytesIO()
    document.save(output)
    return output.getvalue()


def field_notice_html(model: str, revision: str, status: str, lot: str, action: str) -> bytes:
    return f"""<!doctype html><html lang=\"zh-CN\"><body>
<h1>{model} 现场更正通知 Revision {revision}</h1>
<p>状态：{status}</p><p>适用批次：{lot}</p><p>通知编号：C3-FSN-026</p>
<h2>处理要求</h2><p>{action}</p></body></html>""".encode()


def login(client: httpx.Client, api: str) -> str:
    response = client.post(
        f"{api}/api/v1/auth/login",
        json={
            "email": os.getenv("RAGLAB_STACK_SMOKE_EMAIL", "admin@raglab.local"),
            "password": os.getenv("RAGLAB_PLATFORM_ADMIN_PASSWORD", "change-this-admin-password"),
        },
    )
    response.raise_for_status()
    return response.json()["access_token"]


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--api", default=os.getenv("RAGLAB_API_URL", f"http://127.0.0.1:{os.getenv('RAGLAB_API_PORT', '8080')}"))
    parser.add_argument("--dataset", default="public-medical-device")
    parser.add_argument("--max-runes", type=int, default=700)
    parser.add_argument("--overlap-runes", type=int, default=80)
    parser.add_argument("--split", choices=("development", "holdout"), default="development")
    parser.add_argument(
        "--pipeline-release",
        default=os.getenv("DOCUMENT_QUALITY_PIPELINE_RELEASE", ""),
    )
    parser.add_argument("--output", required=True)
    arguments = parser.parse_args()
    font_file = cjk_font()
    with tempfile.TemporaryDirectory(prefix="raglab-document-quality-") as directory:
        root = Path(directory)
        development_sources = {
            ("dev-ocr-aed-critical-001", "synthetic-aed-troubleshooting-r1"): scanned_aed(font_file),
            ("dev-cleaning-margin-002", "synthetic-vsm100-operator-r1"): repeated_margin_pdf(font_file),
            ("dev-table-version-003", "synthetic-vsm-compatibility-r1"): compatibility_xlsx(root / "vsm-compatibility.xlsx"),
            ("dev-chunk-long-procedure-004", "synthetic-long-service-procedure-r1"): long_procedure_markdown(),
            ("dev-ocr-degraded-fallback-005", "synthetic-monitor-degraded-ocr-r1"): degraded_monitor_development_scan(font_file),
            ("dev-version-scope-filter-006", "synthetic-vsm460-network-r4"): b"# VSM-460 5.1 Network Recovery\n\nVSM-460 software 5.1 NET-SYNC-311: export diagnostics, renew the current certificate, verify gateway and time sync, then observe for 15 minutes.\n",
            ("dev-version-scope-filter-006", "synthetic-vsm460-network-r3"): b"# VSM-460 4.7 Historical Recovery\n\nVSM-460 software 4.7 NET-SYNC-311: import the fixed-address backup. Superseded by revision R4 and forbidden for software 5.1.\n",
            ("dev-model-suffix-filter-007", "synthetic-vsm420-pro-power-r1"): power_docx("VSM-420 Pro", "R1", "保留增强电源日志并检查双路冗余输入。"),
            ("dev-model-suffix-filter-007", "synthetic-vsm420-power-r1"): power_docx("VSM-420", "R1", "检查标准电源模块和单路输入。"),
            ("dev-lot-scope-filter-008", "synthetic-c5-field-notice-r2"): field_notice_html("BeneHeart C5", "2", "current", "LOT-M2701", "Quarantine sales stock and install insulation kit C5-M27."),
            ("dev-lot-scope-filter-008", "synthetic-c5-field-notice-r1"): field_notice_html("BeneHeart C5", "1", "superseded", "LOT-M2602", "Historical lot only; do not apply to LOT-M2701."),
        }
        holdout_sources = {
            ("holdout-version-conflict-004", "synthetic-vsm450-network-r3"): vsm450_network_current_markdown(),
            ("holdout-version-conflict-004", "synthetic-vsm450-network-r2"): vsm450_network_history_markdown(),
            ("holdout-similar-model-005", "synthetic-vsm410-pro-power-r1"): power_docx("VSM-410 Pro", "R1", "保留增强电源模块日志并检查双路输入。"),
            ("holdout-similar-model-005", "synthetic-vsm410-power-r1"): power_docx("VSM-410", "R1", "检查标准电源模块和单路输入。"),
            ("holdout-superseded-notice-006", "synthetic-c3-field-notice-r3"): field_notice_html("BeneHeart C3", "3", "当前有效", "LOT-K2608", "停止销售库存并由授权人员安装绝缘垫片套件 FSN-K26。"),
            ("holdout-superseded-notice-006", "synthetic-c3-field-notice-r1"): field_notice_html("BeneHeart C3", "1", "已被 Revision 3 取代", "LOT-K2501", "旧批次执行目视检查；不得用于 LOT-K2608。"),
        }
        cases = DEVELOPMENT_CASES if arguments.split == "development" else HOLDOUT_CASES
        sources = development_sources if arguments.split == "development" else holdout_sources
        artifacts = []
        with httpx.Client(timeout=420) as client:
            token = login(client, arguments.api.rstrip("/"))
            headers = {"Authorization": f"Bearer {token}"}
            for case_id, document_id, filename, content_type in cases:
                response = client.post(
                    f"{arguments.api.rstrip('/')}/api/v1/datasets/{arguments.dataset}/documents/evaluation-artifacts",
                    headers=headers,
                    data={
                        "case_id": case_id,
                        "document_id": document_id,
                        "metadata": json.dumps(DOCUMENT_METADATA.get(document_id, {}), ensure_ascii=False),
                        "max_runes": str(arguments.max_runes),
                        "overlap_runes": str(arguments.overlap_runes),
                    },
                    files={"file": (filename, sources[(case_id, document_id)], content_type)},
                )
                response.raise_for_status()
                artifact = response.json()
                assert artifact["indexed"] is False, artifact
                artifacts.append(artifact)
        output = Path(arguments.output)
        output.parent.mkdir(parents=True, exist_ok=True)
        observed_parsers = sorted({
            f"{item.get('document_ir', {}).get('quality', {}).get('parser', 'unknown')}@"
            f"{item.get('document_ir', {}).get('quality', {}).get('parser_version', 'unknown')}"
            for item in artifacts
        })
        # Release identity must be invariant across corpora. The previous
        # implementation derived it from parsers observed in the current
        # sample, so an OCR-heavy Development split and a native-format
        # Holdout from the same deployment produced different fingerprints.
        pipeline_release = arguments.pipeline_release or "document-ir-v4+chunker-v1+artifact-contract-v2+scope-filter-v1"
        output.write_text(json.dumps({
            "schema": "agent-evaluation.document-quality.artifacts.v2",
            "source": f"rag-evolution-lab authenticated no-index {arguments.split} preview",
            "config": {
                "max_runes": arguments.max_runes,
                "overlap_runes": arguments.overlap_runes,
                "pipeline_release": pipeline_release,
                "observed_parsers": observed_parsers,
            },
            "artifacts": artifacts,
        }, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")
        print(json.dumps({
            "status": "exported",
            "cases": len(artifacts),
            "output": str(output.resolve()),
            "indexed": False,
            "statuses": {item["case_id"]: item["status"] for item in artifacts},
            "split": arguments.split,
        }, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
